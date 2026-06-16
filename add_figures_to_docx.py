"""
add_figures_to_docx.py
======================
Chèn 10 hình minh họa (figures/*.png) vào bản đã trau chuốt
Fair_Article_VN/paper_revised_VN.docx — đúng vị trí từng mục,
giữ nguyên toàn bộ văn bản và bảng.

Chạy: PYTHONUTF8=1 uv run python add_figures_to_docx.py
"""
import os
import shutil
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = "Fair_Article_VN/paper_revised_VN.docx"
FIG = "Fair_Article_VN/figures"

# (anchor_paragraph_index, [(filename, width_inches, caption), ...])
# anchor index theo file gốc; hình được chèn NGAY SAU đoạn anchor.
PLAN = [
    (36, [  # cuối §4.1, trước §4.2
        ("fig_neo4j_schema.png", 5.9,
         "Hình 1a. Sơ đồ schema đồ thị thuộc tính (property graph): năm nhãn node "
         "(tô màu theo vai trò trong SCM) và năm loại quan hệ. Node Post giữ vai trò "
         "trung tâm; Subreddit và User là các biến gây nhiễu (confounder) mà nhánh "
         "nhân quả cô lập."),
        ("fig_graph_stats.png", 6.1,
         "Hình 1b. Quy mô đồ thị HIN Fakeddit được lưu trữ trong Neo4j: số lượng node "
         "và quan hệ theo từng loại."),
    ]),
    (63, [  # §6.1, ngay sau câu dẫn "Hình 2 trực quan hóa..."
        ("fig2_standard_vs_shift.png", 6.2,
         "Hình 2. Độ chính xác OOD trên hai chế độ đánh giá. Khi không có dịch chuyển "
         "gây nhiễu (trái), cả bốn phương pháp tương đương về mặt thống kê (≈57–61%). "
         "Dưới confounder đảo ngược (phải), chúng phân kỳ mạnh: baseline rơi xuống dưới "
         "mức ngẫu nhiên, trong khi CausalHeteroGNN vẫn duy trì độ chính xác cao. "
         "Thanh sai số biểu diễn ±1 độ lệch chuẩn qua 3 seed."),
    ]),
    (78, [  # cuối §6.3, trước §6.4
        ("fig1_confounding_ood_bar.png", 5.8,
         "Hình 3. Độ chính xác OOD dưới confounder đảo ngược (trung bình ± độ lệch chuẩn, "
         "3 seed). Baseline rơi xuống dưới đường ngẫu nhiên 50%; các phương pháp bất biến "
         "mềm (IRM, EERM) hồi phục một phần; structural cut đạt cao nhất và ổn định nhất."),
        ("fig3_confounding_metrics.png", 6.2,
         "Hình 4. Cùng một thứ bậc được bảo toàn trên cả ba độ đo OOD (Accuracy, Macro-F1, "
         "AUC). Giá trị AUC < 0,5 của baseline cho thấy thứ tự xếp hạng bị đảo ngược một "
         "cách chủ động; CausalHeteroGNN đạt AUC 0,85."),
        ("fig7_confusion.png", 6.0,
         "Hình 5. Ma trận nhầm lẫn dưới confounder đảo ngược (n = 298). Baseline gán sai cả "
         "hai lớp (đảo ngược dự đoán), trong khi mô hình nhân quả khôi phục một biên quyết "
         "định dựa trên nội dung."),
    ]),
    (84, [  # cuối §6.4, trước §6.5
        ("fig6_fastrp_leak.png", 5.5,
         "Hình 6. Việc đưa đặc trưng FastRP (tính trên đồ thị đầy đủ) thổi phồng độ chính "
         "xác OOD lên ≈94% cho cả hai nhánh — một hiện tượng rò rỉ thông tin; đặc trưng "
         "chỉ-nội-dung bộc lộ độ khó thực sự ≈57%."),
    ]),
    (91, [  # cuối §6.5, trước §6.6
        ("fig5_lfr.png", 5.9,
         "Hình 7. Tỷ lệ Lật nhãn (Label-Flip-Rate) dưới các can thiệp phẫu thuật đồ thị. "
         "Mô hình nhân quả bất biến tuyệt đối với confounder subreddit (0,0%) nhưng lại "
         "nhạy cảm hơn baseline với các yếu tố nhân quả thực (hình ảnh, domain)."),
    ]),
    (94, [  # cuối §6.6, trước §6.7
        ("fig4_f1_drop.png", 5.6,
         "Hình 8. Mức suy giảm Macro-F1 từ môi trường đồng pha sang môi trường đảo ngược "
         "(càng thấp càng bền). Thứ bậc phản chiếu trật tự độ chính xác: structural cut "
         "suy giảm ít nhất (12,7%), baseline nhiều nhất (63,3%)."),
    ]),
    (105, [  # cuối §6.9, trước §7
        ("fig_bi_overview.png", 6.3,
         "Hình 9. Các phân tích BI tiêu biểu kết xuất từ đồ thị đã lưu: (a) cộng đồng "
         "Louvain trên node Post; (b) các domain có tỷ lệ tin giả cao nhất; (c) tỷ lệ giả "
         "gần 0/1 của từng subreddit — xác nhận trực quan rằng cộng đồng định nghĩa nhãn "
         "(confounder mà nghiên cứu trung hòa); (d) hồ sơ người phát tán theo tỷ lệ giả."),
    ]),
]


def make_image_par(doc, path, width_in):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(FIG, path), width=Inches(width_in))
    return p


def make_caption_par(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p


def main():
    shutil.copy(DOCX, DOCX + ".bak")  # backup
    doc = docx.Document(DOCX)

    # Lấy tham chiếu phần tử anchor TRƯỚC khi chèn (index sẽ dịch sau khi chèn)
    anchors = [(doc.paragraphs[idx]._p, figs) for idx, figs in PLAN]

    n = 0
    for anchor_el, figs in anchors:
        cur = anchor_el
        for fname, w, cap in figs:
            img_p = make_image_par(doc, fname, w)
            cap_p = make_caption_par(doc, cap)
            cur.addnext(img_p._p)      # chèn ảnh ngay sau anchor hiện tại
            img_p._p.addnext(cap_p._p) # chèn caption ngay sau ảnh
            cur = cap_p._p             # hình kế tiếp nối sau caption này
            n += 1

    doc.save(DOCX)
    print(f"Đã chèn {n} hình vào {DOCX} (backup: {DOCX}.bak)")


if __name__ == "__main__":
    main()
