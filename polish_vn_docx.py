# -*- coding: utf-8 -*-
"""
Create a polished Vietnamese version of CausalHeteroGNN_FAIR_VN_Final.docx.

The script keeps tables and embedded figures intact, and only rewrites selected
paragraphs to improve academic tone, coherence, and terminology consistency.
"""

from docx import Document


SRC = "CausalHeteroGNN_FAIR_VN_Final.docx"
OUT = "CausalHeteroGNN_FAIR_VN_Final_trau_chuot.docx"


REPLACEMENTS = {
    1: "Nghiên cứu mô hình đồ thị nhân quả dị thể dựa trên GraphSAGE cho bài toán phát hiện tin giả đa phương thức",
    2: (
        "Abstract — Mạng nơ-ron đồ thị dị thể (Heterogeneous Graph Neural Networks — HGNNs) "
        "là một hướng tiếp cận giàu tiềm năng đối với bài toán phát hiện tin giả đa phương thức, "
        "do có khả năng mô hình hóa đồng thời nội dung, tác giả, nguồn tin, cộng đồng đăng tải "
        "và hình ảnh. Tuy nhiên, khi được huấn luyện trên dữ liệu mạng xã hội, các mô hình đồ thị "
        "dễ khai thác những tương quan giả, đặc biệt là lịch sử nhãn của cộng đồng Subreddit, "
        "thay vì học các tín hiệu nội dung có tính ổn định. Nghiên cứu này đề xuất "
        "CausalHeteroGNN, một kiến trúc Heterogeneous GraphSAGE kết hợp tín hiệu nhất quán "
        "văn bản–hình ảnh (CLIPcons) và được diễn giải dưới khung mô hình nhân quả cấu trúc "
        "(SCM). Điểm cốt lõi của phương pháp là xây dựng một nhánh nhân quả trên đồ thị đã can "
        "thiệp, trong đó các cạnh liên quan đến Subreddit được loại bỏ nhằm hạn chế đường truyền "
        "của biến gây nhiễu. Trên Confounding-Shift Benchmark, CausalHeteroGNN đạt 79.9±4.2% "
        "OOD Accuracy, AUC 0.922 và F1-drop chỉ 5.4%, vượt đáng kể Baseline GNN (ERM, "
        "52.1±7.8%), IRM (54.1±4.7%), MLP content-only (59.5±1.7%) và EERM (59.6±1.1%). "
        "Mô hình cũng đạt Worst-Group Accuracy 37.8±11.3%, cao nhất trong các phương pháp "
        "so sánh. Trên các phép chia cộng đồng tự nhiên LOCO, CausalHeteroGNN đạt 69.4±5.4%, "
        "tương đương EERM (70.0±6.0%) và cao hơn IRM (65.8±2.9%). Các kết quả cho thấy can "
        "thiệp cấu trúc trên đồ thị là một hướng khả thi để nâng cao tính bền vững của mô hình "
        "phát hiện tin giả trong bối cảnh dịch chuyển phân phối."
    ),
    5: (
        "Sự phát triển nhanh của truyền thông đa phương tiện kéo theo sự gia tăng của các hình "
        "thức tin giả ngày càng tinh vi, trong đó thông tin sai lệch có thể được thể hiện đồng "
        "thời qua văn bản, hình ảnh, nguồn phát tán và ngữ cảnh cộng đồng. Bối cảnh này đặt ra "
        "thách thức đáng kể cho các hệ thống phát hiện tự động, bởi mô hình không chỉ cần nhận "
        "diện dấu hiệu bất thường trong nội dung mà còn phải hiểu được quan hệ giữa các thực thể "
        "trong không gian mạng xã hội. Nhờ khả năng biểu diễn nhiều loại nút và quan hệ, mạng "
        "nơ-ron đồ thị dị thể (HGNNs) là một lựa chọn phù hợp cho bài toán phát hiện tin giả đa "
        "phương thức. Tuy vậy, việc bổ sung cấu trúc đồ thị không mặc nhiên bảo đảm khả năng "
        "tổng quát hóa tốt hơn. Trong nhiều trường hợp, GNN có thể học các tương quan giả trong "
        "dữ liệu huấn luyện, đặc biệt là lịch sử phân bố nhãn của cộng đồng đăng tải (Subreddit), "
        "thay vì dựa vào các tín hiệu nội dung ổn định."
    ),
    6: (
        "Dưới góc nhìn nhân quả, Subreddit có thể được xem như một biến gây nhiễu trong cấu trúc "
        "dữ liệu. Với tập Fakeddit [1], một số cộng đồng có tương quan mạnh với nhãn thật/giả, "
        "tạo thành đường đi tắt giữa ngữ cảnh đăng tải và nhãn dự đoán. Khi thông tin này được "
        "lan truyền qua các cạnh của đồ thị, Baseline GNN (ERM HeteroSAGE) có nguy cơ khuếch đại "
        "shortcut cộng đồng và suy giảm đáng kể khi phân phối cộng đồng thay đổi. Ngược lại, "
        "MLP content-only có thể ổn định hơn trong một số phép chia tự nhiên do chỉ dựa vào nội "
        "dung, nhưng lại bỏ qua các quan hệ có giá trị giữa bài đăng, người dùng, nguồn tin và "
        "cộng đồng."
    ),
    7: (
        "Do đó, vấn đề trọng tâm không phải là lựa chọn giữa mô hình nội dung thuần và mô hình "
        "đồ thị, mà là thiết kế một cơ chế khai thác đồ thị có kiểm soát. Nghiên cứu này đề xuất "
        "CausalHeteroGNN, một kiến trúc Heterogeneous GraphSAGE theo hướng nhân quả, trong đó "
        "các đường truyền thông tin liên quan đến Subreddit được loại bỏ khỏi nhánh nhân quả. "
        "Cơ chế này cho phép mô hình tận dụng cấu trúc quan hệ mà không phụ thuộc quá mức vào "
        "shortcut cộng đồng, qua đó duy trì năng lực xử lý nội dung trong các phép chia cộng "
        "đồng tự nhiên và tránh sự suy giảm mạnh khi tương quan cộng đồng–nhãn bị thay đổi hoặc "
        "đảo ngược."
    ),
    8: (
        "Nghiên cứu có ba đóng góp chính. Thứ nhất, bài viết đề xuất CausalHeteroGNN, một kiến "
        "trúc HGNN đa phương thức dựa trên SCM, trong đó Subreddit được mô hình hóa như biến gây "
        "nhiễu và được kiểm soát bằng can thiệp cấu trúc. Thứ hai, nghiên cứu xây dựng một quy "
        "trình đánh giá nhiều tầng gồm Held-Out Subreddit OOD, Confounding-Shift Benchmark, "
        "LOCO trên cộng đồng tự nhiên và các phép can thiệp phản thực tế. Thứ ba, bài viết phân "
        "tích tính bền vững của mô hình thông qua AUC, F1-drop, Worst-Group Accuracy và "
        "Label-Flip Rate, qua đó làm rõ khi nào đồ thị hỗ trợ tổng quát hóa và khi nào đồ thị "
        "trở thành nguồn shortcut."
    ),
    10: (
        "Các nghiên cứu về phát hiện tin giả đa phương thức cho thấy thông tin sai lệch hiếm khi "
        "chỉ nằm trong văn bản; nó còn gắn với hình ảnh, nguồn tin, cộng đồng đăng tải và các "
        "quan hệ xã hội xung quanh bài viết. SAFE [12] khai thác tính nhất quán giữa văn bản và "
        "hình ảnh; BiGCN [11] mô hình hóa quá trình lan truyền; KGAT [14] bổ sung tri thức ngoài; "
        "trong khi HGNNs [4] cho phép biểu diễn nhiều loại thực thể và quan hệ trong cùng một "
        "không gian học. Tuy nhiên, cấu trúc đồ thị cũng có thể mang theo các tín hiệu gây nhiễu. "
        "MLP content-only đôi khi đạt hiệu năng mạnh nhờ tín hiệu văn bản–hình ảnh, trong khi GNN "
        "trên đồ thị thô có nguy cơ học shortcut từ cộng đồng hoặc nguồn phát tán."
    ),
    11: (
        "Từ góc độ suy luận nhân quả, Pearl [2] nhấn mạnh vai trò của can thiệp trong việc kiểm "
        "soát ảnh hưởng của biến gây nhiễu. Schölkopf và Peters [6, 9] cho rằng khả năng tổng "
        "quát hóa phụ thuộc vào việc học được các cơ chế ổn định qua nhiều môi trường, thay vì "
        "chỉ tối ưu các tương quan quan sát được trong dữ liệu huấn luyện. Trong Fakeddit, "
        "Subreddit là một biến gây nhiễu quan trọng vì lịch sử nhãn của cộng đồng thường tương "
        "quan mạnh với nhãn bài đăng. Nếu tín hiệu này được lan truyền qua đồ thị, mô hình có thể "
        "đưa ra dự đoán dựa trên cộng đồng thay vì bằng chứng nội dung."
    ),
    12: (
        "IRM và EERM [13] tiếp cận vấn đề dịch chuyển phân phối bằng các ràng buộc bất biến. IRM "
        "khuyến khích một bộ phân loại ổn định trên nhiều môi trường, còn EERM tạo các môi trường "
        "đồ thị ảo thông qua nhiễu loạn cạnh và giảm phương sai rủi ro giữa chúng. Dù vậy, các "
        "phương pháp này vẫn chủ yếu là ràng buộc mềm và không loại bỏ trực tiếp đường truyền "
        "thông tin từ Subreddit. Khoảng trống này là cơ sở để nghiên cứu đề xuất CausalHeteroGNN, "
        "trong đó can thiệp cấu trúc được sử dụng như một cơ chế kiểm soát shortcut rõ ràng và "
        "có thể diễn giải."
    ),
    13: "3. Data Collection and Representation in Neo4j — Thu thập và biểu diễn dữ liệu trong Neo4j",
    14: (
        "Để quản lý mạng thông tin đa phương thức và dị thể, nghiên cứu sử dụng Neo4j [15] làm "
        "nền tảng lưu trữ, truy vấn và kiểm tra đồ thị. Neo4j hỗ trợ biểu diễn trực quan các "
        "thực thể và quan hệ, đồng thời cho phép kiểm tra tính toàn vẹn dữ liệu bằng Cypher trước "
        "khi chuyển sang giai đoạn huấn luyện mô hình học máy."
    ),
    15: (
        "Nguồn dữ liệu chính là Fakeddit [1], một tập dữ liệu chuẩn quy mô lớn được thu thập từ "
        "Reddit trong giai đoạn 2008–2020, bao gồm văn bản, hình ảnh và nhãn phân loại. Do giới "
        "hạn hạ tầng thực nghiệm, nghiên cứu sử dụng 5.898 bài đăng có ảnh thật. Sau khi mô hình "
        "hóa dưới dạng đồ thị thuộc tính, tập dữ liệu gồm 17.079 nút và 28.274 quan hệ."
    ),
    16: (
        "Dữ liệu được biểu diễn dưới dạng đồ thị thuộc tính trong Neo4j với năm loại nút: Post, "
        "User, Subreddit, Domain và Image; cùng năm loại quan hệ có hướng: POSTED_BY, POSTED_IN, "
        "LINKS_TO, HAS_IMAGE và MEMBER_OF (Hình 1). Cụ thể, đồ thị bao gồm 5.898 nút Post, "
        "4.604 nút User, 658 nút Domain, 21 nút Subreddit và 5.898 nút Image."
    ),
    19: (
        "Để đánh giá khả năng tổng quát hóa ngoài phân phối (OOD), hai cộng đồng r/neutralnews "
        "và r/theonion được giữ lại cho tập kiểm thử OOD. Tập huấn luyện gồm 5.000 bài đăng cân "
        "bằng nhãn (2.500 real và 2.500 fake), tập validation gồm 400 bài, và tập kiểm thử gồm "
        "498 bài, trong đó có 298 mẫu OOD. Sau tiền xử lý, dữ liệu được chuyển thành đối tượng "
        "HeteroData của PyTorch Geometric. Đặc trưng nút Post gồm embedding tiêu đề "
        "(all-mpnet-base-v2), các thuộc tính thống kê (score, upvote ratio, số bình luận) và "
        "CLIPcons, tức chỉ số đo mức nhất quán văn bản–hình ảnh. Đặc trưng hình ảnh được trích "
        "xuất bằng CLIP ViT-B/32; các nút User, Subreddit và Domain được mô tả bằng thống kê "
        "hành vi và lịch sử chỉ tính từ tập huấn luyện nhằm hạn chế rò rỉ nhãn."
    ),
    20: "4. Proposed CausalHeteroGNN with Heterogeneous GraphSAGE",
    21: (
        "CausalHeteroGNN được thiết kế nhằm khai thác ngữ cảnh quan hệ trong đồ thị dị thể nhưng "
        "đồng thời hạn chế sự phụ thuộc vào shortcut cộng đồng. Mô hình sử dụng Heterogeneous "
        "GraphSAGE [3] làm bộ mã hóa nền tảng và gồm hai luồng xử lý song song: nhánh phụ "
        "(spurious branch) nhận đồ thị gốc G với đầy đủ quan hệ, còn nhánh nhân quả (causal "
        "branch) nhận đồ thị đã can thiệp G_causal, trong đó các cạnh chạm tới nút Subreddit bị "
        "loại bỏ. Dự đoán chính của mô hình được lấy từ nhánh nhân quả."
    ),
    23: (
        "Hình 2. Kiến trúc CausalHeteroGNN. Từ HeteroData đầu vào, mô hình tạo hai đồ thị: đồ thị "
        "gốc G và đồ thị nhân quả G_causal bằng cách loại bỏ các cạnh liên quan đến Subreddit. "
        "Hai đồ thị được mã hóa bằng Heterogeneous GraphSAGE dùng chung trọng số để tạo "
        "h_spurious và h_causal. Dự đoán Fake/Real được lấy từ nhánh nhân quả "
        "(h_causal → Classifier), trong khi GRL và ràng buộc trực giao được sử dụng để giảm "
        "thông tin cộng đồng trong h_causal và tách biệt tín hiệu gây nhiễu khỏi tín hiệu nhân quả."
    ),
    24: (
        "Do các loại nút có không gian đặc trưng khác nhau, mỗi loại nút được ánh xạ vào không "
        "gian ẩn chung d = 96 bằng một phép biến đổi tuyến tính riêng, theo sau là ReLU và dropout "
        "(p = 0.4). Hai lớp Heterogeneous GraphSAGE thực hiện lan truyền thông tin theo từng loại "
        "quan hệ r ∈ R; biểu diễn mới của mỗi nút được hình thành bằng cách kết hợp trạng thái "
        "hiện tại với trung bình cộng biểu diễn của các láng giềng liên quan."
    ),
    27: (
        "Để giảm rò rỉ gián tiếp của thông tin cộng đồng, mô hình sử dụng ba cơ chế bổ sung: "
        "(1) h_causal → GRL → Confounder Classifier với loss L_adv, khiến biểu diễn nhân quả khó "
        "dự đoán Subreddit; (2) h_spurious → Confounder Classifier với loss L_spurious, nhằm gom "
        "tín hiệu cộng đồng vào nhánh phụ; và (3) ràng buộc trực giao L_ortho giữa h_spurious và "
        "h_causal, khuyến khích hai nhánh học các không gian biểu diễn khác biệt."
    ),
    31: "5.1 Thiết lập thực nghiệm",
    32: (
        "Nghiên cứu đánh giá CausalHeteroGNN theo ba khía cạnh: (i) khả năng tổng quát hóa ngoài "
        "phân phối (OOD), (ii) khả năng chống shortcut cộng đồng, và (iii) độ tin cậy của quy "
        "trình đánh giá. Giao thức Held-Out Subreddit OOD giữ lại r/neutralnews và r/theonion "
        "cho tập kiểm thử. Giao thức Confounding-Shift Benchmark xây dựng hai cộng đồng tổng hợp "
        "với tương quan nhiễu mạnh ở pha train (ρ = 0.9) và bị đảo ở pha OOD (ρ = 0.1). LOCO "
        "đánh giá trên ba cặp cộng đồng tự nhiên chưa xuất hiện trong huấn luyện. "
        "CausalHeteroGNN được so sánh với bốn đối chứng: MLP content-only, Baseline GNN "
        "(ERM HeteroSAGE), IRM [8] và EERM [13]. Các chỉ số đánh giá gồm Accuracy, Macro-F1, "
        "AUC, F1-drop, Worst-Group Accuracy và Label-Flip Rate (LFR)."
    ),
    34: (
        "Bảng 1 trình bày kết quả chính trên hai giao thức OOD. Trên Held-Out Subreddit, các mô "
        "hình đạt hiệu năng tương đối gần nhau: MLP content-only đạt 60.5±1.1%, EERM đạt "
        "60.9±1.9%, CausalHeteroGNN đạt 59.6±1.9% và Baseline GNN đạt 57.6±2.7%. Kết quả này "
        "cho thấy trong các cộng đồng chưa thấy, tín hiệu nội dung vẫn giữ vai trò trung tâm; "
        "việc bổ sung đồ thị thô không tự động tạo ra lợi thế nếu mô hình không kiểm soát được "
        "các yếu tố gây nhiễu."
    ),
    35: (
        "Trái lại, trên Confounding-Shift, khác biệt giữa các phương pháp trở nên rõ rệt. "
        "Baseline GNN chỉ đạt 52.1±7.8% với AUC 0.511, gần mức ngẫu nhiên, cho thấy mô hình bị "
        "ảnh hưởng mạnh bởi tương quan cộng đồng–nhãn trong pha huấn luyện. IRM (54.1±4.7%) và "
        "EERM (59.6±1.1%) cải thiện một phần nhưng vẫn còn hạn chế. CausalHeteroGNN đạt "
        "79.9±4.2%, AUC 0.922 và F1-drop chỉ 5.4%, chứng tỏ mô hình vừa có khả năng phân biệt "
        "tốt, vừa bền vững hơn trước dịch chuyển của yếu tố gây nhiễu."
    ),
    39: (
        "Hình 3. So sánh Accuracy, AUC và F1-drop trên hai giao thức OOD. CausalHeteroGNN vượt "
        "trội trên Confounding-Shift, đồng thời duy trì hiệu năng cạnh tranh trên "
        "Held-Out Subreddit."
    ),
    41: (
        "Để kiểm tra liệu kết quả trên Confounding-Shift có phụ thuộc vào benchmark tổng hợp hay "
        "không, nghiên cứu tiếp tục đánh giá LOCO trên ba cặp cộng đồng tự nhiên chưa xuất hiện "
        "trong huấn luyện. Với mỗi split, toàn bộ thống kê nút được tính lại chỉ từ tập huấn luyện "
        "nhằm hạn chế rò rỉ nhãn và bảo đảm tính nghiêm ngặt của đánh giá."
    ),
    44: (
        "Kết quả cho thấy Baseline GNN suy giảm ổn định trên mọi fold (56.5±1.3%), qua đó xác "
        "nhận rằng shortcut cộng đồng không chỉ xuất hiện trong benchmark tổng hợp mà còn tồn tại "
        "trong các phép chia tự nhiên. CausalHeteroGNN (69.4±5.4%) và EERM (70.0±6.0%) đạt hiệu "
        "năng tương đương, nhưng EERM có độ dao động lớn hơn. IRM (65.8±2.9%) ổn định hơn nhưng "
        "thấp hơn CausalHeteroGNN. MLP (71.6±7.9%) đạt trung bình cao nhất, song dao động mạnh "
        "theo fold, phản ánh sự phụ thuộc vào đặc trưng nội dung đặc thù của từng cộng đồng."
    ),
    46: (
        "Hình 4. So sánh năm mô hình trên ba fold LOCO tự nhiên và giá trị trung bình. "
        "CausalHeteroGNN và EERM đạt hiệu năng tương đương, đồng thời vượt IRM và Baseline GNN "
        "trên tất cả các fold."
    ),
    48: (
        "Accuracy trung bình có thể che khuất hiệu năng trên các nhóm khó, đặc biệt trong bối "
        "cảnh phân phối bị đảo tương quan. Vì vậy, nghiên cứu tính Worst-Group Accuracy theo "
        "nhóm env × label trên Confounding-Shift."
    ),
    51: (
        "CausalHeteroGNN đạt Worst-Group Accuracy 37.8±11.3%, cao nhất trong các phương pháp so "
        "sánh, đồng thời có khoảng cách Avg−Worst chỉ 33.0 điểm, nhỏ nhất trong nhóm. IRM "
        "(24.4%) và EERM (28.6%) cải thiện nhẹ so với Baseline GNN (23.5%) nhưng vẫn suy giảm "
        "mạnh ở nhóm bị đảo tương quan. Việc CausalHeteroGNN có Avg-Group thấp hơn (70.8%) cho "
        "thấy một đánh đổi có chủ đích giữa tối ưu hiệu năng trung bình và nâng cao tính bền "
        "vững trên nhóm cực trị."
    ),
    53: (
        "Nghiên cứu thực hiện các phép can thiệp do(·) trên đồ thị nhân quả và đo tỷ lệ thay đổi "
        "dự đoán nhãn (Label-Flip Rate — LFR). Chỉ số này cho phép đánh giá mức độ nhạy của mô "
        "hình trước các thay đổi có kiểm soát trong cấu trúc dữ liệu."
    ),
    56: (
        "Baseline GNN phản ứng mạnh với phép hoán đổi cộng đồng (30.1%), trong khi nhánh nhân quả "
        "gần như bất biến (~0.0%), phù hợp với thiết kế loại bỏ các cạnh liên quan đến Subreddit. "
        "Đổi lại, nhánh nhân quả nhạy hơn với can thiệp trên hình ảnh (4.0% → 8.0%) và nguồn tin "
        "(6.2% → 15.5%). Điều này cho thấy sau khi giảm phụ thuộc vào cộng đồng, mô hình chuyển "
        "trọng tâm sang các tín hiệu nội dung và nguồn tin có ý nghĩa hơn đối với nhiệm vụ phân "
        "loại."
    ),
    58: (
        "Hình 5. Phân tích robustness: (a) quan hệ giữa Worst-Group Accuracy và Avg-Group Accuracy "
        "trên Confounding-Shift; (b) Label-Flip Rate dưới các phép can thiệp cấu trúc do(·)."
    ),
    60: (
        "Hình 6 tổng hợp các kết quả thực nghiệm dưới dạng BI dashboard. Sáu thẻ KPI chính gồm: "
        "Conf-Shift OOD Accuracy (79.9±4.2%), AUC (0.922), F1-drop (5.4%), Worst-Group Accuracy "
        "(37.8±11.3%), LFR do(Subreddit) (~0.0%) và LOCO Mean Accuracy (69.4±5.4%). Các panel "
        "bên dưới trình bày chi tiết từng khía cạnh, bao gồm hiệu năng trên Confounding-Shift, "
        "AUC và F1-drop, xu hướng LOCO, Worst-Group Accuracy, Label-Flip Rate và biểu đồ tổng "
        "hợp KPI của ba mô hình tiêu biểu."
    ),
    63: (
        "Dashboard cho thấy CausalHeteroGNN là mô hình cân bằng tốt nhất giữa các tiêu chí đánh "
        "giá: OOD Accuracy cao trên Confounding-Shift, AUC vượt trội, F1-drop thấp, "
        "Worst-Group Accuracy cao nhất và LFR gần như bằng 0 khi can thiệp vào Subreddit, trong "
        "khi vẫn duy trì hiệu năng LOCO tương đương EERM."
    ),
    65: (
        "Nghiên cứu đã đề xuất CausalHeteroGNN, một mô hình Heterogeneous GraphSAGE theo hướng "
        "nhân quả cho bài toán phát hiện tin giả đa phương thức. Bằng cách xác định Subreddit là "
        "biến gây nhiễu và loại bỏ các cạnh chạm tới Subreddit trong nhánh nhân quả, mô hình hạn "
        "chế việc học các tương quan giả từ cộng đồng đăng tải. Dự đoán chính được lấy trực tiếp "
        "từ h_causal → Classifier; GRL chỉ được sử dụng như một cơ chế phụ trợ để làm giảm thông "
        "tin Subreddit trong biểu diễn nhân quả."
    ),
    66: (
        "Kết quả trên Confounding-Shift Benchmark cho thấy CausalHeteroGNN (CLIPcons) đạt "
        "79.9±4.2% OOD Accuracy và AUC 0.922, vượt Baseline GNN ERM (52.1±7.8%), IRM "
        "(54.1±4.7%), MLP content-only (59.5±1.7%) và EERM (59.6±1.1%). Mô hình đạt "
        "Worst-Group Accuracy 37.8±11.3%, cao nhất trong các phương pháp so sánh, trong khi "
        "F1-drop chỉ 5.4%; ngược lại, Baseline GNN suy giảm tới 45.1%. Trên LOCO tự nhiên, "
        "CausalHeteroGNN đạt 69.4±5.4%, cao hơn Baseline GNN (56.5±1.3%) và IRM (65.8±2.9%), "
        "đồng thời tương đương EERM (70.0±6.0%) nhưng ổn định hơn."
    ),
    67: (
        "Những kết quả này củng cố nhận định rằng can thiệp cấu trúc kết hợp với tín hiệu "
        "CLIPcons là một hướng tiếp cận khả thi để kiểm soát shortcut cộng đồng trong phát hiện "
        "tin giả đa phương thức. Trong các nghiên cứu tiếp theo, phương pháp có thể được mở rộng "
        "trên tập dữ liệu quy mô lớn hơn, đánh giá trong bối cảnh phân loại đa lớp, đồng thời "
        "tích hợp dashboard BI thành công cụ hỗ trợ phân tích, giám sát và kiểm toán mô hình."
    ),
}


def set_paragraph_text(paragraph, text):
    """Replace paragraph text while keeping the first run's formatting."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def main():
    doc = Document(SRC)
    for index, text in REPLACEMENTS.items():
        set_paragraph_text(doc.paragraphs[index], text)
    doc.save(OUT)
    print(f"Saved polished document: {OUT}")


if __name__ == "__main__":
    main()
