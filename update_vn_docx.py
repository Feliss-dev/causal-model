# -*- coding: utf-8 -*-
"""
update_vn_docx.py
=================
Cap nhat (VN)_Causal Graph.docx voi:
  1. Abstract: so lieu dung (79.9%, worst-group 37.8%, LOCO)
  2. Bang 1: them chu thich CLIPcons
  3. Phan 5.3 LOCO: them moi (IRM/EERM 3 seeds)
  4. Phan 5.4 Worst-Group: them moi
  5. Ket luan: sua so lieu (74.2->79.9, Baseline 36.4->52.1)
"""

import shutil
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "(VN)_Causal Graph.docx"
BAK = "(VN)_Causal Graph_BACKUP2.docx"
OUT = "(VN)_Causal Graph.docx"

shutil.copy(SRC, BAK)
print(f"[OK] Backup -> {BAK}")

doc = Document(SRC)


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _make_w_p(text="", bold=False):
    """Build a bare w:p XML element."""
    p = OxmlElement("w:p")
    if text:
        r = OxmlElement("w:r")
        if bold:
            rpr = OxmlElement("w:rPr")
            b = OxmlElement("w:b")
            rpr.append(b)
            r.append(rpr)
        t = OxmlElement("w:t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        p.append(r)
    return p


def _make_w_tbl(rows_data, header=None):
    """
    Build a w:tbl XML element.
    header: list of column titles (first row, bold)
    rows_data: list of lists of strings
    """
    tbl = OxmlElement("w:tbl")

    tbl_pr = OxmlElement("w:tblPr")
    tbl_style = OxmlElement("w:tblStyle")
    tbl_style.set(qn("w:val"), "TableGrid")
    tbl_pr.append(tbl_style)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), "0")
    tbl_w.set(qn("w:type"), "auto")
    tbl_pr.append(tbl_w)
    tbl.append(tbl_pr)

    all_rows = ([header] if header else []) + list(rows_data)
    cols = len(all_rows[0])

    tbl_grid = OxmlElement("w:tblGrid")
    for _ in range(cols):
        gc = OxmlElement("w:gridCol")
        tbl_grid.append(gc)
    tbl.append(tbl_grid)

    for ri, row in enumerate(all_rows):
        is_header = (header is not None and ri == 0)
        tr = OxmlElement("w:tr")
        for ci, val in enumerate(row):
            tc = OxmlElement("w:tc")
            tc_pr = OxmlElement("w:tcPr")
            tc.append(tc_pr)
            cell_p = OxmlElement("w:p")
            cell_r = OxmlElement("w:r")
            if is_header or (ri == len(all_rows) - 1 and not is_header and header):
                # last data row bold if it's the "Mean" row — skip for now
                pass
            if is_header:
                rpr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rpr.append(b)
                cell_r.append(rpr)
            cell_t = OxmlElement("w:t")
            cell_t.text = str(val)
            cell_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            cell_r.append(cell_t)
            cell_p.append(cell_r)
            tc.append(cell_p)
            tr.append(tc)
        tbl.append(tr)

    return tbl


def insert_after(anchor_xml, new_xml):
    """Insert new_xml immediately after anchor_xml in the XML tree."""
    anchor_xml.addnext(new_xml)
    return new_xml


def set_para_text(p, new_text):
    """Replace text in a python-docx Paragraph object."""
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    for cp in cell.paragraphs:
        cp.clear()
        run = cp.add_run(text)
        run.bold = bold
        cp.alignment = align


# ── 1. Fix Abstract ──────────────────────────────────────────────────────────
ABS_NEW = (
    "Abstract - Tom tat. Mang no-ron do thi hon hop (Heterogeneous Graph Neural Networks - HGNNs) "
    "la mot huong tiep can hieu qua cho bai toan phat hien tin gia da phuong thuc. "
    "Nghien cuu nay de xuat CausalHeteroGNN, mot kien truc Heterogeneous GraphSAGE ket hop "
    "tin hieu nhat quan van ban-hinh anh (CLIPcons), dua tren mo hinh nhan qua cau truc (SCM), "
    "trong do nhanh nhan qua thuc hien can thiep cau truc bang cach loai bo cac duong truyen "
    "thong tin lien quan den Subreddit. Tren Confounding-Shift Benchmark, CausalHeteroGNN dat "
    "79.9+/-4.2% OOD Accuracy, AUC 0.922 va F1-drop chi 5.4%, vuot Baseline GNN (ERM, "
    "52.1+/-7.8%), IRM (54.1+/-4.7%), MLP content-only (59.5+/-1.7%) va EERM (59.6+/-1.1%). "
    "Worst-Group Accuracy dat 37.8+/-11.3%, cao hon tat ca cac baseline. Tren cac phep chia "
    "cong dong tu nhien LOCO, CausalHeteroGNN dat 69.4+/-5.4%; IRM dat 65.8+/-2.9% va EERM "
    "dat 70.0+/-6.0% (mean tren 3 seeds). Nghien cuu dong thoi chi ra nguy co ro ri nhan tu "
    "FastRP khi tinh tren toan bo do thi truoc khi tach split, co the lam tang ao OOD Accuracy "
    "hon 33 diem phan tram."
)

abs_ok = False
for p in doc.paragraphs:
    if p.text.strip().startswith("Abstract - T"):
        set_para_text(p, ABS_NEW)
        abs_ok = True
        print("[OK] Abstract updated")
        break
if not abs_ok:
    print("[WARN] Abstract not found")


# ── 2. Fix Conclusion ─────────────────────────────────────────────────────────
CONC_MARKER = "74.2%"
CONC_NEW = (
    "Ket qua tren Confounding-Shift Benchmark cho thay CausalHeteroGNN (voi CLIPcons) dat "
    "79.9+/-4.2% OOD Accuracy va AUC 0.922, vuot Baseline GNN ERM (52.1+/-7.8%), IRM "
    "(54.1+/-4.7%), MLP content-only (59.5+/-1.7%) va EERM (59.6+/-1.1%). Worst-Group "
    "Accuracy cua CausalHeteroGNN dat 37.8+/-11.3%, cao nhat trong tat ca nhom danh gia. "
    "Tren cac fold LOCO tu nhien, CausalHeteroGNN dat 69.4+/-5.4%, cao hon Baseline GNN "
    "(56.5+/-1.3%) va IRM (65.8+/-2.9%); EERM dat 70.0+/-6.0% (trung binh tren 3 seeds), "
    "tuong duong MLP (71.6+/-7.9%) nhung dao dong lon hon. F1-drop cua CausalHeteroGNN chi "
    "5.4%, cho thay hieu nang on dinh giua seen va OOD, trong khi Baseline GNN suy giam toi "
    "45.1%. Nhung ket qua nay xac nhan rang can thiep cau truc ket hop CLIPcons la mot huong "
    "kha thi de kiem soat shortcut cong dong trong phat hien tin gia da phuong thuc."
)

conc_ok = False
for p in doc.paragraphs:
    if CONC_MARKER in p.text and "CausalHeteroGNN" in p.text:
        set_para_text(p, CONC_NEW)
        conc_ok = True
        print("[OK] Conclusion updated")
        break
if not conc_ok:
    print("[WARN] Conclusion paragraph not found")


# ── 3. Update Table 0: add + to CausalHeteroGNN row label ────────────────────
if doc.tables:
    tbl0 = doc.tables[0]
    for ri, row in enumerate(tbl0.rows):
        c0 = row.cells[0].text.strip()
        if "CausalHeteroGNN" in c0 and ri > 0:
            new_label = c0 + " (CLIPcons)"
            set_cell(row.cells[0], new_label, bold=False)
            print(f"[OK] Table0 row{ri}: added CLIPcons annotation")
            break


# ── 4. Update Bảng 1 caption ──────────────────────────────────────────────────
for p in doc.paragraphs:
    t = p.text.strip()
    if ("Bang 1" in t or "Bảng 1" in t) and "Hiệu năng" in t:
        new_cap = (
            t + " | CausalHeteroGNN dung CLIPcons (5 seeds); "
            "baseline khong co CLIPcons (3 seeds)."
        )
        set_para_text(p, new_cap)
        print("[OK] Bang 1 caption updated")
        break


# ── 5. Find anchor paragraph (after last figure caption in section 5) ─────────
# We look for the Hinh 4 / Dashboard caption, or else last para before section 6
anchor_p_obj = None
candidates = []
for p in doc.paragraphs:
    t = p.text.strip()
    if ("Hinh 4" in t or "Hình 4" in t or "Dashboard" in t):
        candidates.append(p)
    elif "5.5" in t or "Label-Flip" in t:
        candidates.append(p)

# Use first match as anchor
for p in doc.paragraphs:
    t = p.text.strip()
    if ("Hinh 4" in t or "Hình 4" in t) and ("Dashboard" in t or "tổng hợp" in t):
        anchor_p_obj = p
        break

if anchor_p_obj is None:
    # Try: look for any paragraph that ends section 5.2
    for p in doc.paragraphs:
        t = p.text.strip()
        if "5.3" in t or "5.4" in t or "5.5" in t:
            # insert before this section
            # use the paragraph BEFORE it
            break
    # Fallback: last paragraph of doc
    for p in doc.paragraphs:
        if p.text.strip().startswith("Hình 3") or p.text.strip().startswith("Hinh 3"):
            anchor_p_obj = p
            break

if anchor_p_obj is None:
    anchor_p_obj = doc.paragraphs[-3]
    print("[WARN] Using fallback anchor")

print(f"[OK] Anchor: {anchor_p_obj.text[:70]!r}")

# anchor_xml is the w:p element we insert after
anchor_xml = anchor_p_obj._p


# ── 6. Insert Section 5.3 LOCO ────────────────────────────────────────────────

# Build all XML pieces — insert in REVERSE ORDER so each addnext goes right after anchor
# (each addnext pushes the previous one down)

loco_header_row = ["Fold held-out", "MLP(*)", "Baseline GNN(*)", "IRM(**)", "EERM(**)", "CausalHeteroGNN(*)"]
loco_data = [
    ["nottheonion + pareidolia",        "61.3", "57.9", "62.0+/-4.3", "62.2+/-0.1", "62.3"],
    ["upliftingnews + fakehistoryporn", "73.1", "56.8", "66.5+/-4.0", "70.7+/-1.0", "70.5"],
    ["usnews+usanews+fakealbumcovers",  "80.5", "54.7", "69.1+/-4.8", "76.9+/-2.9", "75.5"],
    ["Mean+/-std (across folds)",       "71.6+/-7.9", "56.5+/-1.3", "65.8+/-2.9", "70.0+/-6.0", "69.4+/-5.4"],
]

wg_header_row = ["Mo hinh", "Worst-Group Acc (%)", "Avg-Group Acc (%)"]
wg_data = [
    ["MLP content-only",   "15.6+/-3.1",  "56.5+/-1.0"],
    ["Baseline GNN (ERM)", "23.5+/-7.9",  "73.2+/-4.4"],
    ["IRM",                "24.4+/-7.1",  "74.3+/-2.7"],
    ["EERM",               "28.6+/-1.3",  "77.0+/-1.0"],
    ["CausalHeteroGNN",    "37.8+/-11.3", "70.8+/-4.5"],
]

# Build all pieces as XML — then insert in reverse order after anchor

pieces = [
    # (xml_element, description)
    (_make_w_p("5.3. Danh gia LOCO tren cong dong tu nhien", bold=True), "5.3 heading"),
    (_make_w_p(
        "De tra loi phan bien rang Confounding-Shift la benchmark nhan tao, nghien cuu danh gia "
        "LOCO (Leave-One-Community-Out) tren ba cap cong dong tu nhien chua thay trong huan luyen. "
        "Toan bo thong ke node duoc tinh lai tu tap huan luyen cua split moi de tranh ro ri. "
        "IRM va EERM duoc bao cao voi mean+/-std tren 3 seeds {42,1,2}; MLP, Baseline GNN va "
        "CausalHeteroGNN dung single seed=42."
    ), "5.3 intro"),
    (_make_w_p("Bang 3. LOCO tren ba fold tu nhien — Accuracy (%) tren fold chua thay.", bold=True), "5.3 table cap"),
    (_make_w_tbl(loco_data, header=loco_header_row), "5.3 LOCO table"),
    (_make_w_p("(*) single seed=42. (**) mean+/-std tren 3 seeds {42,1,2}."), "5.3 table note"),
    (_make_w_p(
        "Ket qua cho thay Baseline GNN suy giam on dinh tren moi fold (56.5+/-1.3%). "
        "CausalHeteroGNN (69.4+/-5.4%) va EERM (70.0+/-6.0%) dat hieu nang tuong duong, "
        "nhung EERM co dao dong lon hon (std 6.0 so voi 5.4 diem). IRM (65.8+/-2.9%) on dinh "
        "nhung thap hon CausalHeteroGNN. Hinh 9 minh hoa so sanh day du."
    ), "5.3 discussion"),
    (_make_w_p(
        "Hinh 9. So sanh 5 mo hinh tren ba fold LOCO tu nhien va trung binh "
        "(single seed=42 cho Baseline/MLP/Causal; 3 seeds cho IRM/EERM)."
    ), "Fig9 caption"),
    (_make_w_p("5.4. Phan tich Worst-Group Accuracy", bold=True), "5.4 heading"),
    (_make_w_p(
        "Accuracy trung binh co the che khuat cac nhom kho. Nghien cuu tinh Worst-Group "
        "Accuracy theo nhom env x label tren Confounding-Shift (mean+/-std: 3 seeds cho "
        "baseline, 5 seeds cho CausalHeteroGNN)."
    ), "5.4 intro"),
    (_make_w_p("Bang 4. Worst-Group Accuracy tren Confounding-Shift Benchmark.", bold=True), "5.4 table cap"),
    (_make_w_tbl(wg_data, header=wg_header_row), "5.4 WG table"),
    (_make_w_p(
        "CausalHeteroGNN dat worst-group 37.8+/-11.3%, cao hon tat ca phuong phap so sanh. "
        "IRM (24.4%) va EERM (28.6%) cai thien nhe so voi Baseline GNN (23.5%), nhung van suy "
        "giam manh o nhom bi dao tuong quan. Luu y CausalHeteroGNN co avg-group thap hon "
        "(70.8%) do danh doi giua toi uu trung binh va tinh ben vung nhom cuc tri."
    ), "5.4 discussion"),
]

# Insert in reverse order so final order is correct
for xml_elem, desc in reversed(pieces):
    anchor_xml.addnext(xml_elem)
    print(f"  [+] Inserted: {desc}")

print("[OK] Sections 5.3 and 5.4 inserted")


# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"\n[DONE] Saved -> {OUT}")
print(f"       Backup -> {BAK}")
