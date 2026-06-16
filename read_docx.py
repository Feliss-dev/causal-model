import docx
doc = docx.Document("(VN)_Causal Graph.docx")
print("=== TABLES ===")
for ti, tbl in enumerate(doc.tables):
    print(f"\n--- Table {ti} ({len(tbl.rows)} rows x {len(tbl.columns)} cols) ---")
    for ri, row in enumerate(tbl.rows):
        cells = [c.text.strip()[:35] for c in row.cells]
        sep = " | "
        print(f"  row{ri}: " + sep.join(cells))

print("\n=== ALL PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(f"[{i:03d}] " + repr(t[:130]))
