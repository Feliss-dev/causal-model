# -*- coding: utf-8 -*-
"""
Create a clearer academic Vietnamese version of the CausalHeteroGNN paper.

This version keeps the scientific content, numbers, tables, and figures, but
uses shorter sentences and explains technical claims in a more reader-friendly
style.
"""

from docx import Document


SRC = "CausalHeteroGNN_FAIR_VN_Final_trau_chuot_sua_hinh4.docx"
OUT = "CausalHeteroGNN_FAIR_VN_Final_de_hieu.docx"


REPLACEMENTS = {
    1: "Mô hình đồ thị nhân quả dị thể dựa trên GraphSAGE cho phát hiện tin giả đa phương thức",
    2: (
        "Abstract — Phát hiện tin giả đa phương thức là bài toán khó vì thông tin sai lệch có "
        "thể xuất hiện đồng thời trong văn bản, hình ảnh, nguồn đăng tải và cộng đồng lan truyền. "
        "Mạng nơ-ron đồ thị dị thể (Heterogeneous Graph Neural Networks — HGNNs) phù hợp với bài "
        "toán này vì có thể biểu diễn nhiều loại thực thể như bài đăng, người dùng, hình ảnh, "
        "nguồn tin và Subreddit trong cùng một đồ thị. Tuy nhiên, mô hình đồ thị cũng có thể học "
        "các đường tắt trong dữ liệu. Ví dụ, nếu một Subreddit thường chứa nhiều tin giả trong tập "
        "huấn luyện, mô hình có thể dựa quá nhiều vào cộng đồng đó thay vì phân tích nội dung bài "
        "đăng. Nghiên cứu này đề xuất CausalHeteroGNN, một mô hình Heterogeneous GraphSAGE có "
        "thêm nhánh nhân quả. Trong nhánh này, các cạnh liên quan đến Subreddit được loại bỏ để "
        "giảm ảnh hưởng của biến gây nhiễu. Mô hình cũng sử dụng tín hiệu CLIPcons để đo mức nhất "
        "quán giữa văn bản và hình ảnh. Trên Confounding-Shift Benchmark, CausalHeteroGNN đạt "
        "79.9±4.2% OOD Accuracy, AUC 0.922 và F1-drop 5.4%, cao hơn rõ rệt so với Baseline GNN "
        "(52.1±7.8%), IRM (54.1±4.7%), MLP content-only (59.5±1.7%) và EERM (59.6±1.1%). "
        "Mô hình cũng đạt Worst-Group Accuracy 37.8±11.3%, cao nhất trong các phương pháp so "
        "sánh. Kết quả cho thấy can thiệp cấu trúc trên đồ thị là một hướng hiệu quả để giảm "
        "shortcut cộng đồng và cải thiện khả năng tổng quát hóa."
    ),
    5: (
        "Sự phát triển của mạng xã hội và truyền thông đa phương tiện làm cho tin giả ngày càng "
        "khó phát hiện. Một bài đăng sai lệch không chỉ thể hiện qua nội dung văn bản, mà còn có "
        "thể liên quan đến hình ảnh, người đăng, nguồn tin và cộng đồng nơi bài viết xuất hiện. "
        "Vì vậy, các hệ thống phát hiện tin giả cần khai thác nhiều nguồn thông tin cùng lúc. "
        "Mạng nơ-ron đồ thị dị thể (HGNNs) là một hướng tiếp cận phù hợp vì có thể mô hình hóa "
        "nhiều loại nút và quan hệ trong cùng một cấu trúc đồ thị. Tuy nhiên, đồ thị không phải "
        "lúc nào cũng giúp mô hình tổng quát hóa tốt hơn. Nếu không kiểm soát cẩn thận, GNN có "
        "thể học các tương quan giả trong tập huấn luyện, đặc biệt là mối liên hệ giữa Subreddit "
        "và nhãn thật/giả."
    ),
    6: (
        "Trong nghiên cứu này, Subreddit được xem là một biến gây nhiễu. Lý do là một số cộng "
        "đồng trong Fakeddit [1] có xu hướng gắn với một loại nhãn nhất định. Khi mô hình nhìn "
        "thấy một bài đăng thuộc cộng đồng đó, nó có thể dự đoán nhãn dựa trên lịch sử của cộng "
        "đồng thay vì dựa trên nội dung thật sự của bài viết. Đây là một dạng shortcut. Shortcut "
        "này có thể giúp mô hình đạt kết quả cao trên dữ liệu giống tập huấn luyện, nhưng dễ làm "
        "mô hình suy giảm khi gặp cộng đồng mới hoặc khi phân phối dữ liệu thay đổi."
    ),
    7: (
        "Vấn đề đặt ra là cần khai thác lợi ích của đồ thị nhưng không để mô hình phụ thuộc quá "
        "mức vào shortcut cộng đồng. Để giải quyết vấn đề này, nghiên cứu đề xuất "
        "CausalHeteroGNN. Mô hình vẫn sử dụng Heterogeneous GraphSAGE để học quan hệ giữa các "
        "thực thể, nhưng tạo thêm một nhánh nhân quả. Trong nhánh này, các cạnh liên quan đến "
        "Subreddit được loại bỏ. Nhờ đó, mô hình buộc phải dựa nhiều hơn vào nội dung, hình ảnh, "
        "nguồn tin và các quan hệ ít bị nhiễu hơn."
    ),
    8: (
        "Nghiên cứu có ba đóng góp chính. Thứ nhất, bài viết đề xuất CausalHeteroGNN, một mô "
        "hình HGNN đa phương thức có cơ chế can thiệp cấu trúc để giảm ảnh hưởng của Subreddit. "
        "Thứ hai, nghiên cứu đánh giá mô hình bằng nhiều kịch bản khác nhau, gồm Held-Out "
        "Subreddit OOD, Confounding-Shift Benchmark, LOCO trên cộng đồng tự nhiên và các phép "
        "can thiệp phản thực tế. Thứ ba, bài viết phân tích độ bền vững của mô hình qua nhiều "
        "chỉ số như AUC, F1-drop, Worst-Group Accuracy và Label-Flip Rate."
    ),
    10: (
        "Các nghiên cứu trước đây cho thấy phát hiện tin giả đa phương thức cần kết hợp nhiều "
        "loại tín hiệu. SAFE [12] tập trung vào sự nhất quán giữa văn bản và hình ảnh. BiGCN "
        "[11] khai thác quá trình lan truyền thông tin. KGAT [14] bổ sung tri thức ngoài. Các "
        "mô hình HGNNs [4] cho phép kết hợp nhiều loại thực thể và quan hệ trong cùng một đồ thị. "
        "Tuy nhiên, cấu trúc đồ thị cũng có thể chứa thông tin gây nhiễu. Vì vậy, dùng đồ thị "
        "một cách trực tiếp chưa chắc đã tốt hơn mô hình chỉ dùng nội dung."
    ),
    11: (
        "Theo lý thuyết nhân quả, một mô hình muốn tổng quát hóa tốt cần học các cơ chế ổn định, "
        "không chỉ học các tương quan xuất hiện trong tập huấn luyện [2, 6, 9]. Trong Fakeddit, "
        "Subreddit là một ví dụ rõ ràng của biến gây nhiễu. Một số cộng đồng có lịch sử nhãn lệch "
        "về tin thật hoặc tin giả. Nếu GNN lan truyền thông tin này qua đồ thị, mô hình có thể "
        "dự đoán dựa trên cộng đồng thay vì dựa trên bằng chứng nội dung."
    ),
    12: (
        "IRM và EERM [13] là hai hướng tiếp cận nhằm cải thiện khả năng tổng quát hóa khi dữ liệu "
        "bị dịch chuyển. IRM khuyến khích mô hình học bộ phân loại ổn định qua nhiều môi trường. "
        "EERM tạo thêm các môi trường đồ thị ảo bằng cách nhiễu loạn cạnh. Tuy nhiên, các phương "
        "pháp này không trực tiếp loại bỏ đường truyền thông tin từ Subreddit. CausalHeteroGNN "
        "khác ở chỗ can thiệp thẳng vào cấu trúc đồ thị bằng cách cắt các cạnh liên quan đến "
        "Subreddit trong nhánh nhân quả."
    ),
    14: (
        "Nghiên cứu sử dụng Neo4j [15] để lưu trữ và kiểm tra đồ thị dữ liệu. Neo4j phù hợp với "
        "bài toán này vì có thể biểu diễn rõ các loại thực thể khác nhau và các quan hệ giữa "
        "chúng. Trước khi đưa dữ liệu vào mô hình học máy, các truy vấn Cypher được dùng để kiểm "
        "tra tính đầy đủ và nhất quán của đồ thị."
    ),
    15: (
        "Dữ liệu chính là Fakeddit [1], một tập dữ liệu lớn được thu thập từ Reddit trong giai "
        "đoạn 2008–2020. Tập dữ liệu này gồm văn bản, hình ảnh và nhãn phân loại. Do giới hạn hạ "
        "tầng thực nghiệm, nghiên cứu sử dụng 5.898 bài đăng có ảnh thật. Sau khi chuyển sang "
        "dạng đồ thị, dữ liệu gồm 17.079 nút và 28.274 quan hệ."
    ),
    16: (
        "Đồ thị gồm năm loại nút: Post, User, Subreddit, Domain và Image. Các quan hệ có hướng "
        "gồm POSTED_BY, POSTED_IN, LINKS_TO, HAS_IMAGE và MEMBER_OF (Hình 1). Cụ thể, đồ thị có "
        "5.898 nút Post, 4.604 nút User, 658 nút Domain, 21 nút Subreddit và 5.898 nút Image."
    ),
    19: (
        "Để đánh giá khả năng tổng quát hóa ngoài phân phối (OOD), nghiên cứu giữ lại hai cộng "
        "đồng r/neutralnews và r/theonion cho tập kiểm thử OOD. Tập huấn luyện gồm 5.000 bài "
        "đăng cân bằng nhãn, tập validation gồm 400 bài, và tập kiểm thử gồm 498 bài. Trong đó, "
        "298 mẫu thuộc nhóm OOD. Sau tiền xử lý, dữ liệu được chuyển thành HeteroData của "
        "PyTorch Geometric. Đặc trưng của Post gồm embedding tiêu đề, các thống kê như score, "
        "upvote ratio, số bình luận, và CLIPcons. CLIPcons đo mức độ phù hợp giữa văn bản và "
        "hình ảnh. Đặc trưng ảnh được trích xuất bằng CLIP ViT-B/32. Các đặc trưng của User, "
        "Subreddit và Domain chỉ được tính từ tập huấn luyện để hạn chế rò rỉ nhãn."
    ),
    20: "4. Mô hình CausalHeteroGNN với Heterogeneous GraphSAGE",
    21: (
        "CausalHeteroGNN được thiết kế để tận dụng thông tin quan hệ trong đồ thị nhưng giảm phụ "
        "thuộc vào Subreddit. Mô hình dùng Heterogeneous GraphSAGE [3] làm bộ mã hóa chính và "
        "gồm hai nhánh. Nhánh phụ nhận đồ thị gốc G với đầy đủ quan hệ. Nhánh nhân quả nhận đồ "
        "thị G_causal, trong đó các cạnh liên quan đến Subreddit đã bị loại bỏ. Kết quả dự đoán "
        "Fake/Real được lấy từ nhánh nhân quả."
    ),
    23: (
        "Hình 2. Kiến trúc CausalHeteroGNN. Từ HeteroData đầu vào, mô hình tạo hai đồ thị: đồ "
        "thị gốc G và đồ thị nhân quả G_causal. G_causal được tạo bằng cách loại bỏ các cạnh "
        "liên quan đến Subreddit. Hai đồ thị được mã hóa bằng Heterogeneous GraphSAGE dùng chung "
        "trọng số để tạo h_spurious và h_causal. Dự đoán Fake/Real được lấy từ h_causal. GRL và "
        "ràng buộc trực giao được dùng để giảm thông tin cộng đồng trong biểu diễn nhân quả."
    ),
    24: (
        "Vì mỗi loại nút có đặc trưng khác nhau, mô hình trước hết đưa tất cả đặc trưng về cùng "
        "một không gian ẩn có kích thước d = 96. Sau đó, hai lớp Heterogeneous GraphSAGE lan "
        "truyền thông tin theo từng loại quan hệ. Nói cách khác, mỗi nút cập nhật biểu diễn của "
        "mình bằng cách kết hợp thông tin hiện tại với thông tin từ các nút láng giềng."
    ),
    27: (
        "Để hạn chế thông tin Subreddit lọt vào nhánh nhân quả, mô hình dùng thêm ba thành phần. "
        "Thứ nhất, GRL làm cho h_causal khó dự đoán Subreddit. Thứ hai, nhánh phụ h_spurious được "
        "khuyến khích học thông tin cộng đồng. Thứ ba, ràng buộc trực giao làm cho h_causal và "
        "h_spurious khác nhau hơn. Nhờ đó, mô hình tách tín hiệu nội dung ổn định khỏi tín hiệu "
        "cộng đồng dễ gây nhiễu."
    ),
    32: (
        "Nghiên cứu đánh giá CausalHeteroGNN theo ba khía cạnh: khả năng tổng quát hóa ngoài "
        "phân phối, khả năng chống shortcut cộng đồng, và độ tin cậy của quy trình đánh giá. "
        "Held-Out Subreddit OOD giữ lại r/neutralnews và r/theonion cho tập kiểm thử. "
        "Confounding-Shift Benchmark tạo tình huống trong đó tương quan giữa cộng đồng và nhãn "
        "mạnh ở tập huấn luyện nhưng bị đảo ở tập OOD. LOCO kiểm tra mô hình trên các cộng đồng "
        "tự nhiên chưa xuất hiện trong huấn luyện. Các mô hình so sánh gồm MLP content-only, "
        "Baseline GNN, IRM và EERM. Các chỉ số gồm Accuracy, Macro-F1, AUC, F1-drop, "
        "Worst-Group Accuracy và Label-Flip Rate."
    ),
    34: (
        "Bảng 1 trình bày kết quả trên hai giao thức OOD. Trên Held-Out Subreddit, các mô hình "
        "có kết quả khá gần nhau. MLP content-only đạt 60.5±1.1%, EERM đạt 60.9±1.9%, "
        "CausalHeteroGNN đạt 59.6±1.9%, và Baseline GNN đạt 57.6±2.7%. Điều này cho thấy khi "
        "gặp cộng đồng chưa thấy, tín hiệu nội dung vẫn rất quan trọng. Đồ thị chỉ tạo ra lợi "
        "thế khi mô hình kiểm soát được các nguồn gây nhiễu trong đồ thị."
    ),
    35: (
        "Trên Confounding-Shift, sự khác biệt rõ hơn. Baseline GNN chỉ đạt 52.1±7.8% và AUC "
        "0.511, gần mức ngẫu nhiên. Điều này cho thấy mô hình bị ảnh hưởng mạnh bởi shortcut "
        "cộng đồng. IRM và EERM cải thiện một phần nhưng chưa giải quyết triệt để vấn đề. "
        "CausalHeteroGNN đạt 79.9±4.2%, AUC 0.922 và F1-drop chỉ 5.4%. Kết quả này cho thấy "
        "mô hình phân biệt tốt hơn và ổn định hơn khi tương quan cộng đồng–nhãn bị thay đổi."
    ),
    39: (
        "Hình 3. So sánh Accuracy, AUC và F1-drop trên hai giao thức OOD. CausalHeteroGNN nổi "
        "bật trên Confounding-Shift và vẫn giữ hiệu năng cạnh tranh trên Held-Out Subreddit."
    ),
    41: (
        "Để kiểm tra kết quả trong bối cảnh tự nhiên hơn, nghiên cứu đánh giá LOCO trên ba cặp "
        "cộng đồng chưa xuất hiện trong huấn luyện. Với mỗi split, các thống kê của nút chỉ được "
        "tính từ tập huấn luyện. Cách làm này giúp tránh việc thông tin từ tập kiểm thử rò rỉ vào "
        "quá trình huấn luyện."
    ),
    44: (
        "Kết quả LOCO cho thấy Baseline GNN giảm hiệu năng khá ổn định trên cả ba fold "
        "(56.5±1.3%). Điều này cho thấy shortcut cộng đồng không chỉ xuất hiện trong benchmark "
        "tổng hợp mà còn tồn tại trong dữ liệu tự nhiên. CausalHeteroGNN đạt 69.4±5.4%, gần với "
        "EERM (70.0±6.0%) và cao hơn IRM (65.8±2.9%). MLP đạt trung bình cao nhất "
        "(71.6±7.9%) nhưng dao động lớn hơn, cho thấy mô hình này phụ thuộc khá mạnh vào đặc "
        "trưng nội dung của từng cộng đồng."
    ),
    46: (
        "Hình 4. So sánh năm mô hình trên ba fold LOCO tự nhiên và giá trị trung bình. "
        "CausalHeteroGNN và EERM đạt hiệu năng tương đương, đồng thời cao hơn IRM và Baseline "
        "GNN trên tất cả các fold."
    ),
    48: (
        "Accuracy trung bình không phải lúc nào cũng phản ánh đầy đủ chất lượng của mô hình. "
        "Một mô hình có thể đạt điểm trung bình cao nhưng vẫn dự đoán kém trên một số nhóm dữ "
        "liệu khó, nhất là khi mối quan hệ giữa cộng đồng và nhãn bị thay đổi. Vì vậy, nghiên "
        "cứu sử dụng Worst-Group Accuracy để đo hiệu năng thấp nhất của mô hình trên các nhóm "
        "env × label trong Confounding-Shift."
    ),
    51: (
        "CausalHeteroGNN đạt Worst-Group Accuracy 37.8±11.3%, cao nhất trong các phương pháp so "
        "sánh. Khoảng cách giữa Avg-Group và Worst-Group của mô hình là 33.0 điểm, nhỏ nhất "
        "trong nhóm. Điều này cho thấy CausalHeteroGNN không chỉ tốt ở trung bình mà còn ổn định "
        "hơn trên nhóm khó. IRM và EERM cải thiện nhẹ so với Baseline GNN nhưng vẫn suy giảm "
        "mạnh ở nhóm bị đảo tương quan. Avg-Group của CausalHeteroGNN thấp hơn một số mô hình "
        "khác, cho thấy có sự đánh đổi giữa tối ưu điểm trung bình và bảo vệ nhóm khó."
    ),
    53: (
        "Nghiên cứu tiếp tục dùng các phép can thiệp do(·) để kiểm tra mô hình phụ thuộc vào loại "
        "thông tin nào. Sau mỗi can thiệp, nghiên cứu đo Label-Flip Rate (LFR), tức tỷ lệ mẫu bị "
        "đổi nhãn dự đoán. Nếu LFR cao, mô hình nhạy với loại thông tin vừa bị can thiệp."
    ),
    56: (
        "Baseline GNN rất nhạy với phép hoán đổi cộng đồng, với LFR 30.1%. Ngược lại, nhánh nhân "
        "quả của CausalHeteroGNN gần như không đổi nhãn khi can thiệp vào Subreddit (~0.0%). "
        "Kết quả này phù hợp với thiết kế của mô hình, vì các cạnh liên quan đến Subreddit đã bị "
        "loại bỏ trong nhánh nhân quả. Khi can thiệp vào hình ảnh và nguồn tin, LFR của "
        "CausalHeteroGNN cao hơn. Điều này cho thấy sau khi giảm phụ thuộc vào cộng đồng, mô "
        "hình chuyển sang sử dụng nhiều hơn các tín hiệu nội dung và nguồn tin."
    ),
    58: (
        "Hình 5. Phân tích độ bền vững của mô hình: (a) quan hệ giữa Worst-Group Accuracy và "
        "Avg-Group Accuracy trên Confounding-Shift; (b) Label-Flip Rate dưới các phép can thiệp "
        "cấu trúc do(·)."
    ),
    60: (
        "Hình 6 tổng hợp các kết quả thực nghiệm dưới dạng dashboard. Các chỉ số chính gồm "
        "Conf-Shift OOD Accuracy (79.9±4.2%), AUC (0.922), F1-drop (5.4%), Worst-Group Accuracy "
        "(37.8±11.3%), LFR do(Subreddit) (~0.0%) và LOCO Mean Accuracy (69.4±5.4%). Dashboard "
        "giúp quan sát nhanh hiệu năng của mô hình ở nhiều góc độ thay vì chỉ nhìn vào một chỉ "
        "số duy nhất."
    ),
    63: (
        "Dashboard cho thấy CausalHeteroGNN cân bằng tốt giữa nhiều tiêu chí: Accuracy cao trên "
        "Confounding-Shift, AUC tốt, F1-drop thấp, Worst-Group Accuracy cao nhất và gần như không "
        "đổi nhãn khi can thiệp vào Subreddit. Đồng thời, mô hình vẫn giữ hiệu năng LOCO tương "
        "đương EERM."
    ),
    65: (
        "Nghiên cứu đã đề xuất CausalHeteroGNN, một mô hình Heterogeneous GraphSAGE theo hướng "
        "nhân quả cho phát hiện tin giả đa phương thức. Ý tưởng chính là xem Subreddit như một "
        "biến gây nhiễu và loại bỏ các cạnh liên quan đến Subreddit trong nhánh nhân quả. Nhờ đó, "
        "mô hình giảm khả năng học shortcut cộng đồng. Dự đoán chính được lấy trực tiếp từ "
        "h_causal, còn GRL chỉ đóng vai trò hỗ trợ làm giảm thông tin Subreddit trong biểu diễn."
    ),
    66: (
        "Kết quả thực nghiệm cho thấy CausalHeteroGNN đạt 79.9±4.2% OOD Accuracy và AUC 0.922 "
        "trên Confounding-Shift Benchmark, cao hơn Baseline GNN, IRM, MLP content-only và EERM. "
        "Mô hình cũng đạt Worst-Group Accuracy cao nhất, 37.8±11.3%, và F1-drop chỉ 5.4%. Trên "
        "LOCO tự nhiên, CausalHeteroGNN đạt 69.4±5.4%, cao hơn Baseline GNN và IRM, đồng thời "
        "tương đương EERM nhưng ổn định hơn."
    ),
    67: (
        "Nhìn chung, kết quả cho thấy can thiệp cấu trúc kết hợp với tín hiệu CLIPcons là một "
        "hướng tiếp cận hiệu quả để giảm shortcut cộng đồng trong phát hiện tin giả đa phương "
        "thức. Trong tương lai, phương pháp có thể được mở rộng trên tập dữ liệu lớn hơn, đánh "
        "giá với bài toán phân loại đa lớp, và tích hợp dashboard BI để hỗ trợ phân tích, giám "
        "sát và kiểm toán mô hình."
    ),
}


def set_paragraph_text(paragraph, text):
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def main():
    doc = Document(SRC)
    for index, text in REPLACEMENTS.items():
        set_paragraph_text(doc.paragraphs[index], text)
    doc.save(OUT)
    print(f"Saved easier academic document: {OUT}")


if __name__ == "__main__":
    main()
